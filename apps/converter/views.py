from django.shortcuts import render, redirect
from django.urls import reverse
from django.conf import settings
from django.views import View
# from .services.pdf_service import PDFService
# from .services.word_service import WordService
import os
import random
import string
import sys
import uuid
from django.utils import timezone
from django.http import JsonResponse
from .models import UploadedFile, ProcessedFile, DailyStat, ShortLink
from utils.file_cleanup import cleanup_old_files, cleanup_expired_links_db
from .views_unlock import UnlockPDFView
import fitz
from PIL import Image
import zipfile
import io

class IndexView(View):
    def get(self, request):
        return render(request, 'converter/index.html')

class MergePDFView(View):
    def get(self, request):
        context = {
            'title': 'รวมไฟล์ PDF',
            'subtitle': 'รวมไฟล์ PDF ตามลำดับที่คุณต้องการด้วยตัวรวม PDF ที่ง่ายที่สุด',
            'action_url': '',  # To be implemented
            'file_type': 'PDF',
            'accept': '.pdf'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        files = request.FILES.getlist('files')
        if files:
            uploaded_ids = []
            for f in files:
                uploaded_file = UploadedFile.objects.create(
                    file=f,
                    original_filename=f.name
                )
                uploaded_ids.append(str(uploaded_file.id))
            
            # Save IDs to session
            if request.GET.get('append') == 'true':
                 existing_ids = request.session.get('merge_pdf_file_ids', [])
                 existing_ids.extend(uploaded_ids)
                 request.session['merge_pdf_file_ids'] = existing_ids
                 request.session.modified = True
            else:
                 request.session['merge_pdf_file_ids'] = uploaded_ids
                 request.session.modified = True
            
            return redirect('converter:arrange_pdf')
            
        return redirect('converter:merge_pdf')

class SplitPDFView(View):
    def get(self, request):
        context = {
            'title': 'แยกไฟล์ PDF',
            'subtitle': 'แยกหน้าหนึ่งหน้าหรือทั้งชุดเพื่อแปลงเป็นไฟล์ PDF อิสระได้อย่างง่ายดาย',
            'action_url': reverse('converter:split_pdf'),
            'file_type': 'PDF',
            'accept': '.pdf'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        try:
             # Handle multiple files input (name="files") from template
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # Save uploaded file
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # Prepare output directory
            job_id = str(uuid.uuid4())
            output_dir = os.path.join(settings.MEDIA_ROOT, 'processed', job_id)
            os.makedirs(output_dir, exist_ok=True)
            
            # Split PDF using PyPDF2
            import PyPDF2
            import zipfile
            
            extracted_files = []
            filename_base = os.path.splitext(os.path.basename(input_path))[0]
            
            with open(input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                
                for i in range(total_pages):
                    writer = PyPDF2.PdfWriter()
                    writer.add_page(reader.pages[i])
                    
                    page_filename = f"{filename_base}_page_{i+1}.pdf"
                    page_path = os.path.join(output_dir, page_filename)
                    
                    with open(page_path, 'wb') as out_f:
                        writer.write(out_f)
                    extracted_files.append(page_filename)
            
            # Zip all extracted files
            zip_filename = f"{filename_base}_split.zip"
            zip_path = os.path.join(output_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file in extracted_files:
                    file_abs_path = os.path.join(output_dir, file)
                    zipf.write(file_abs_path, arcname=file)
                    
            # Cleanup: Delete individual PDF pages after zipping
            # ensures only the ZIP file remains for the download view to find
            for file in extracted_files:
                try:
                    os.remove(os.path.join(output_dir, file))
                except Exception as cleanup_error:
                    print(f"Cleanup error: {cleanup_error}")
            
            # Create ProcessedFile record for the ZIP
            processed_rel_path = f"processed/{job_id}/{zip_filename}"
            processed_file = ProcessedFile(
                file=processed_rel_path
            )
            processed_file.save()
            
            # Track usage statistics
            try:
                stat, _ = DailyStat.objects.get_or_create(date=timezone.now().date())
                stat.usage_count += 1
                stat.save()
            except Exception as e:
                print(f"Stats error: {e}")
            
            # Generate ID-only download link
            download_url = reverse('converter:download_file', kwargs={'job_id': job_id})
            
            context = {
                'success': True,
                'download_url': download_url,
                'file_name': zip_filename, # User will download the ZIP
                'file_type': 'ZIP', # For result template
                'file_size': os.path.getsize(zip_path)
            }
            return render(request, 'converter/result.html', context)

        except Exception as e:
            print(f"Error splitting PDF: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาดในการแยกไฟล์: {str(e)}",
                'title': 'แยกไฟล์ PDF'
            }
            return render(request, 'converter/tool_base.html', context)

class PDFToWordView(View):
    def get(self, request):
        context = {
            'title': 'PDF เป็น Word',
            'subtitle': 'แปลงไฟล์ PDF ของคุณเป็นเอกสาร DOC และ DOCX ที่แก้ไขได้ง่ายอย่างง่ายดาย',
            'action_url': reverse('converter:pdf_to_word'),
            'file_type': 'PDF',
            'accept': '.pdf'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        from django.conf import settings
        
        print("DEBUG: PDFToWordView POST Request Received!")
        try:
            # Handle multiple files input (name="files") from template
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0] # Process first file
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # Save uploaded file locally first
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # Worker Configuration
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # Send to Worker
            print(f"📡 Dispatching PDF->Word to Worker: {WORKER_URL}")
            with open(input_path, 'rb') as f:
                # PDF mime type
                mime = 'application/pdf'
                files = {'file': (uploaded_file.name, f, mime)}
                headers = {'X-API-KEY': API_KEY}
                data = {'type': 'pdf-to-word'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    if task_id:
                        return render(request, 'converter/worker_wait.html', {
                            'task_id': task_id,
                            'worker_host': 'https://blilnkdex.biz.id',
                            'file_name': uploaded_file.name,
                            'task_type': 'WORD' # Expected output type
                        })
                    else:
                        raise Exception("Worker did not return Task ID")
                else:
                    raise Exception(f"Worker Error: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"Error converting PDF to Word: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'title': 'PDF เป็น Word'
            }
            return render(request, 'converter/tool_base.html', context)

class WordToPDFView(View):
    def get(self, request):
        context = {
            'title': 'Word เป็น PDF',
            'subtitle': 'แปลงเอกสาร Word ของคุณเป็น PDF ได้อย่างง่ายดาย',
            'action_url': reverse('converter:word_to_pdf'),
            'file_type': 'WORD',
            'accept': '.doc,.docx'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        from django.conf import settings
        
        # DEBUG: Confirm execution of New Logic
        print("🚀 Executing WordToPDFView POST (Worker Node Logic)")
        
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # 1. Save locally
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # 2. Config
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # 3. Dispatch to Worker
            print(f"📡 Dispatching Word->PDF to Worker: {WORKER_URL}")
            with open(input_path, 'rb') as f:
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                files = {'file': (uploaded_file.name, f, mime)}
                headers = {'X-API-KEY': API_KEY}
                data = {'type': 'word-to-pdf'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    if task_id:
                        return render(request, 'converter/worker_wait.html', {
                            'task_id': task_id,
                            'worker_host': 'https://blilnkdex.biz.id',
                            'file_name': uploaded_file.name,
                            'task_type': 'PDF'
                        })
                    else:
                        raise Exception("Worker did not return Task ID")
                else:
                    raise Exception(f"Worker Error: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"Error converting Word to PDF: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'title': 'Word เป็น PDF'
            }
            return render(request, 'converter/tool_base.html', context)

# Direct File Download View (ID Only - Logic resolves filename)
from django.http import FileResponse, Http404
def download_file(request, job_id):
    import os
    from django.conf import settings
    
    # Path to the job directory
    job_dir = os.path.join(settings.MEDIA_ROOT, 'processed', job_id)
    
    print(f"DEBUG: Attempting download for Job ID: {job_id}")
    
    if os.path.exists(job_dir) and os.path.isdir(job_dir):
        # Find the first file in this directory
        files = os.listdir(job_dir)
        if files:
            filename = files[0]
            file_path = os.path.join(job_dir, filename)
            print(f"DEBUG: Found file: {file_path}")
            
            try:
                response = FileResponse(open(file_path, 'rb'), as_attachment=True)
                # Manually set filename header to handle UTF-8 correctly
                from urllib.parse import quote
                response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"
                return response
            except Exception as e:
                print(f"DEBUG: Error opening file: {e}")
                raise Http404(f"Error reading file: {e}")
                
    print("DEBUG: File or directory not found")
    raise Http404("File not found on server")

class CompressPDFView(View):
    def get(self, request):
        context = {
            'title': 'บีบอัด PDF',
            'subtitle': 'ลดขนาดไฟล์ PDF ของคุณโดยไม่สูญเสียคุณภาพมากเกินไป',
            'action_url': '',
            'file_type': 'PDF',
            'accept': '.pdf',
            'show_quality_options': True  # Flag to show compression quality options
        }
        return render(request, 'converter/compress_pdf.html', context)

    def post(self, request):
        files = request.FILES.getlist('files')
        if not files:
            return redirect('converter:compress_pdf')
        
        # Get compression quality from form
        quality = request.POST.get('quality', 'medium')
        
        uploaded_file = UploadedFile.objects.create(
            file=files[0],
            original_filename=files[0].name
        )
        
        try:
            import fitz  # PyMuPDF
            import os
            from django.conf import settings
            
            input_path = uploaded_file.file.path
            output_filename = f"compressed_{uuid.uuid4()}.pdf"
            output_rel_path = f"processed/{output_filename}"
            output_full_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
            
            os.makedirs(os.path.dirname(output_full_path), exist_ok=True)
            
            # Quality settings for PyMuPDF
            quality_settings = {
                'low': {'deflate': 9, 'garbage': 4, 'image_quality': 50},
                'medium': {'deflate': 5, 'garbage': 3, 'image_quality': 75},
                'high': {'deflate': 3, 'garbage': 2, 'image_quality': 90},
                'maximum': {'deflate': 1, 'garbage': 1, 'image_quality': 95}
            }
            
            settings_dict = quality_settings.get(quality, quality_settings['medium'])
            
            # Open PDF
            doc = fitz.open(input_path)
            
            # Compress images in PDF
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get all images on the page
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    
                    # Extract image
                    try:
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Compress image using PIL
                        from PIL import Image
                        import io
                        
                        # Open image
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        
                        # Convert to RGB if necessary
                        if pil_image.mode in ('RGBA', 'LA', 'P'):
                            pil_image = pil_image.convert('RGB')
                        
                        # Compress
                        output_buffer = io.BytesIO()
                        pil_image.save(
                            output_buffer,
                            format='JPEG',
                            quality=settings_dict['image_quality'],
                            optimize=True
                        )
                        
                        # Replace image in PDF
                        compressed_image = output_buffer.getvalue()
                        
                        # Update the existing stream with the compressed image
                        doc.update_stream(xref, compressed_image)
                        
                    except Exception as e:
                        print(f"Error compressing image {img_index} on page {page_num}: {e}")
                        continue
            
            # Save with compression
            doc.save(
                output_full_path,
                garbage=settings_dict['garbage'],
                deflate=True,
                clean=True
            )
            doc.close()
            
            # Get file sizes for comparison
            original_size = os.path.getsize(input_path)
            compressed_size = os.path.getsize(output_full_path)
            
            if compressed_size >= original_size:
                # If compressed file is larger, use original
                import shutil
                shutil.copy(input_path, output_full_path)
                compressed_size = original_size
                reduction_percent = 0
            else:
                reduction_percent = ((original_size - compressed_size) / original_size) * 100
            
            print(f"Original size: {original_size} bytes")
            print(f"Compressed size: {compressed_size} bytes")
            print(f"Reduction: {reduction_percent:.1f}%")
            
            # Save processed file
            processed_file = ProcessedFile(file=output_rel_path)
            processed_file.save()
            
            # Track usage statistics
            try:
                stat, _ = DailyStat.objects.get_or_create(date=timezone.now().date())
                stat.usage_count += 1
                stat.save()
            except Exception as e:
                print(f"Stats error: {e}")
            
            return render(request, 'converter/result.html', {
                'download_url': processed_file.file.url,
                'file_id': processed_file.id,
                'file_type': 'PDF',
                'original_size': original_size,
                'compressed_size': compressed_size,
                'reduction_percent': f"{reduction_percent:.1f}"
            })
            
        except Exception as e:
            print(f"PDF compression error: {e}")
            import traceback
            traceback.print_exc()
            from django.contrib import messages
            messages.error(request, f'เกิดข้อผิดพลาดในการบีบอัดไฟล์: {str(e)}')
            return redirect('converter:compress_pdf')


class PDFToPowerPointView(View):
    def get(self, request):
        from django.urls import reverse
        context = {
            'title': 'PDF เป็น PowerPoint',
            'subtitle': 'เปลี่ยนไฟล์ PDF ของคุณเป็นสไลด์โชว์ PPT และ PPTX ที่แก้ไขได้ง่าย',
            'action_url': reverse('converter:pdf_to_powerpoint'),
            'file_type': 'PDF',
            'accept': '.pdf'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        from django.conf import settings
        
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # Save upload locally first
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # Worker Configuration
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # Send to Worker
            print(f"📡 Dispatching PDF->PPT to Worker: {WORKER_URL}")
            with open(input_path, 'rb') as f:
                files = {'file': (uploaded_file.name, f, 'application/pdf')}
                headers = {'X-API-KEY': API_KEY}
                
                # Send Task Type
                data = {'type': 'pdf-to-ppt'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    if task_id:
                        return render(request, 'converter/worker_wait.html', {
                            'task_id': task_id,
                            'worker_host': 'https://blilnkdex.biz.id',
                            'file_name': uploaded_file.name,
                            'task_type': 'PowerPoint'
                        })
                    else:
                        raise Exception("Worker did not return Task ID")
                else:
                    raise Exception(f"Worker Error: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"Error converting PDF to PPTX: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'title': 'PDF เป็น PowerPoint'
            }
            return render(request, 'converter/tool_base.html', context)


class PDFToExcelView(View):
    def get(self, request):
        from django.urls import reverse
        context = {
            'title': 'PDF เป็น Excel',
            'subtitle': 'ดึงข้อมูลโดยตรงจาก PDF ลงในสเปรดชีต Excel ในเวลาไม่กี่วินาที',
            'action_url': reverse('converter:pdf_to_excel'),
            'file_type': 'PDF',
            'accept': '.pdf'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        from django.conf import settings
        
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # Save upload locally first
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # Worker Configuration
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # Send to Worker
            print(f"📡 Dispatching PDF->Excel to Worker: {WORKER_URL}")
            with open(input_path, 'rb') as f:
                files = {'file': (uploaded_file.name, f, 'application/pdf')}
                headers = {'X-API-KEY': API_KEY}
                
                # Send Task Type
                data = {'type': 'pdf-to-excel'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    if task_id:
                        return render(request, 'converter/worker_wait.html', {
                            'task_id': task_id,
                            'worker_host': 'https://blilnkdex.biz.id',
                            'file_name': uploaded_file.name,
                            'task_type': 'Excel'
                        })
                    else:
                        raise Exception("Worker did not return Task ID")
                else:
                    raise Exception(f"Worker Error: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"Error converting PDF to Excel: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'title': 'PDF เป็น Excel'
            }
            return render(request, 'converter/tool_base.html', context)


class PowerPointToPDFView(View):
    def get(self, request):
        from django.urls import reverse
        context = {
            'title': 'PowerPoint เป็น PDF',
            'subtitle': 'ทำให้สไลด์โชว์ PPT และ PPTX ดูง่ายโดยการแปลงเป็น PDF',
            'action_url': reverse('converter:powerpoint_to_pdf'),
            'file_type': 'POWERPOINT',
            'accept': '.ppt,.pptx'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        from django.conf import settings
        
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # Save upload locally first
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # Worker Configuration
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # Send to Worker
            print(f"📡 Dispatching PPT->PDF to Worker: {WORKER_URL}")
            with open(input_path, 'rb') as f:
                # Note: 'file' field is what api.php expects
                files = {'file': (uploaded_file.name, f, 'application/vnd.openxmlformats-officedocument.presentationml.presentation')}
                headers = {'X-API-KEY': API_KEY}
                
                # Send Task Type
                data = {'type': 'ppt-to-pdf'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    if task_id:
                        return render(request, 'converter/worker_wait.html', {
                            'task_id': task_id,
                            'worker_host': 'https://blilnkdex.biz.id',
                            'file_name': uploaded_file.name,
                            'task_type': 'PDF'  # Target Format
                        })
                    else:
                        raise Exception("Worker did not return Task ID")
                else:
                    raise Exception(f"Worker Error: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"Error converting PPTX to PDF: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'title': 'PowerPoint เป็น PDF'
            }
            return render(request, 'converter/tool_base.html', context)


class WordToPDFView(View):
    def get(self, request):
        context = {
            'title': 'Word เป็น PDF',
            'subtitle': 'ทำให้ไฟล์ DOC และ DOCX อ่านง่ายโดยการแปลงเป็น PDF',
            'action_url': '',
            'file_type': 'WORD',
            'accept': '.doc,.docx'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        from django.conf import settings
        
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            # Save upload locally first
            upload_instance = UploadedFile(file=uploaded_file)
            upload_instance.save()
            input_path = upload_instance.file.path
            
            # Worker Configuration
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # Send to Worker
            print(f"📡 Dispatching Word->PDF to Worker: {WORKER_URL}")
            with open(input_path, 'rb') as f:
                # Word files (doc/docx) mime type
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                files = {'file': (uploaded_file.name, f, mime)}
                headers = {'X-API-KEY': API_KEY}
                data = {'type': 'word-to-pdf'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    if task_id:
                        return render(request, 'converter/worker_wait.html', {
                            'task_id': task_id,
                            'worker_host': 'https://blilnkdex.biz.id',
                            'file_name': uploaded_file.name,
                            'task_type': 'PDF'
                        })
                    else:
                        raise Exception("Worker did not return Task ID")
                else:
                    raise Exception(f"Worker Error: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"Error converting Word to PDF: {e}")
            context = {
                'error': f"เกิดข้อผิดพลาด: {str(e)}",
                'title': 'Word เป็น PDF'
            }
            return render(request, 'converter/tool_base.html', context)

class MergeWordView(View):
    def get(self, request):
        context = {
            'title': 'รวมไฟล์ Word',
            'subtitle': 'รวมไฟล์ Docs และ Docx หลายไฟล์ให้เป็นไฟล์ Word เดียว',
            'action_url': '',
            'file_type': 'WORD',
            'accept': '.doc,.docx'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        files = request.FILES.getlist('files')
        if files:
            uploaded_ids = []
            for f in files:
                uploaded_file = UploadedFile.objects.create(
                    file=f,
                    original_filename=f.name
                )
                uploaded_ids.append(str(uploaded_file.id))
            
            # Save IDs to session to use in Arrange Page
            if request.GET.get('append') == 'true':
                 existing_ids = request.session.get('merge_word_file_ids', [])
                 existing_ids.extend(uploaded_ids)
                 request.session['merge_word_file_ids'] = existing_ids
                 request.session.modified = True
            else:
                 request.session['merge_word_file_ids'] = uploaded_ids
                 request.session.modified = True
            
            return redirect('converter:arrange_word')
            
        return redirect('converter:merge_word')

class ArrangeWordView(View):
    def get(self, request):
        file_ids = request.session.get('merge_word_file_ids', [])
        if not file_ids:
            return redirect('converter:merge_word')
        
        # Preserve order? No, user will order them here.
        # But we need to fetch objects.
        files = UploadedFile.objects.filter(id__in=file_ids)
        
        # Sort files by the order they were in the list if possible, 
        # but filter() querysets are not ordered by input list.
        # We can re-sort them in python to match input order if desired, 
        # or just list them.
        files_dict = {str(f.id): f for f in files}
        ordered_files = [files_dict[fid] for fid in file_ids if fid in files_dict]

        # Extract preview text for each file
        from docx import Document
        for f in ordered_files:
            try:
                doc = Document(f.file.path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                    if len(full_text) >= 5: # Get first 5 paragraphs
                        break
                f.preview_text = full_text
            except Exception as e:
                print(f"Error reading doc {f.id} path {f.file.path}: {e}")
                f.preview_text = []

        print(f"DEBUG: Processed {len(ordered_files)} files.")
        
        # Define alignment mapping
        # 0=LEFT, 1=CENTER, 2=RIGHT, 3=JUSTIFY
        # Note: python-docx alignment might be None if inherited.
        
        for f in ordered_files:
             # Pre-calculate display name
             f.display_name = f.original_filename if f.original_filename else f.filename
             
             # Enhanced Preview Extraction
             try:
                 doc = Document(f.file.path)
                 preview_data = []
                 for para in doc.paragraphs:
                     text = para.text.strip()
                     if text:
                         # Determine Alignment
                         align = 'left'
                         try:
                             # WD_ALIGN_PARAGRAPH.CENTER is 1
                             if para.alignment == 1:
                                 align = 'center'
                             elif para.alignment == 2:
                                 align = 'right'
                         except:
                             pass
                         
                         # Determine Style (Basic)
                         is_bold = False
                         style_name = para.style.name.lower()
                         if 'title' in style_name or 'heading' in style_name:
                             is_bold = True
                             
                         preview_data.append({
                             'text': text,
                             'align': align,
                             'bold': is_bold
                         })
                         
                     if len(preview_data) >= 6:
                         break
                 f.preview_data = preview_data
             except Exception as e:
                 print(f"Preview error for {f.id}: {e}")
                 f.preview_data = []
                 
             print(f"File {f.id} preview lines: {len(getattr(f, 'preview_data', []))}")

        context = {
            'files': ordered_files
        }
        return render(request, 'converter/arrange_word.html', context)

    def post(self, request):
        ordered_ids = request.POST.getlist('file_order[]')
        
        if not ordered_ids:
            return redirect('converter:arrange_word')

        files = []
        try:
            for fid in ordered_ids:
                files.append(UploadedFile.objects.get(id=fid))
        except UploadedFile.DoesNotExist:
            # Handle error
            return redirect('converter:merge_word')

        # Perform Merge
        from .services.word_service import WordService
        service = WordService()
        
        # Prepare file paths
        import os
        from django.conf import settings
        
        file_paths = [f.file.path for f in files]
        
        # Output filename
        output_filename = f"merged_{uuid.uuid4()}.docx"
        output_rel_path = f"processed/{output_filename}"
        output_full_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        
        # Ensure processed dir exists
        os.makedirs(os.path.dirname(output_full_path), exist_ok=True)
        
        try:
            service.merge_word_files(file_paths, output_full_path)
            
            # Create ProcessedFile entry
            # Django FileField expects a path relative to MEDIA_ROOT or a File object.
            # Since we wrote directly to filesystem, we can just save the relative path string
            # IF upload_to handles it, but FileField usually needs a File wrapper to "save".
            # However, for an existing file, we can manually set the name attribute.
            
            processed_file = ProcessedFile(file=output_rel_path)
            processed_file.save()
            
            # Track usage statistics
            try:
                stat, _ = DailyStat.objects.get_or_create(date=timezone.now().date())
                stat.usage_count += 1
                stat.save()
            except Exception as e:
                print(f"Stats error: {e}")
            
            # Cleanup session
            if 'merge_word_file_ids' in request.session:
                del request.session['merge_word_file_ids']
                
            # Redirect to result with success
            # You might want to pass the download link or ID
            return render(request, 'converter/result.html', {
                'download_url': processed_file.file.url,
                'file_id': processed_file.id,
                'file_type': 'WORD'
            })
            
        except Exception as e:
            print(f"Merge Error: {e}")
            # In production show a user friendly error
            return redirect('converter:merge_word')

class ArrangePDFView(View):
    def get(self, request):
        file_ids = request.session.get('merge_pdf_file_ids', [])
        if not file_ids:
            return redirect('converter:merge_pdf')
        
        files = UploadedFile.objects.filter(id__in=file_ids)
        files_dict = {str(f.id): f for f in files}
        ordered_files = [files_dict[fid] for fid in file_ids if fid in files_dict]

        # Extract preview text for PDF
        import PyPDF2
        for f in ordered_files:
            f.display_name = f.original_filename if f.original_filename else f.filename
            try:
                reader = PyPDF2.PdfReader(f.file.path)
                if len(reader.pages) > 0:
                    page = reader.pages[0]
                    text = page.extract_text()
                    lines = [l.strip() for l in text.split('\n') if l.strip()]
                    f.preview_data = [{'text': line, 'align': 'left', 'bold': False} for line in lines[:6]]
                else:
                    f.preview_data = []
            except Exception as e:
                print(f"PDF Preview error {f.id}: {e}")
                f.preview_data = []

        context = {
            'files': ordered_files
        }
        return render(request, 'converter/arrange_pdf.html', context)

    def post(self, request):
        ordered_ids = request.POST.getlist('file_order[]')
        
        if not ordered_ids:
            return redirect('converter:arrange_pdf')

        files = []
        try:
            for fid in ordered_ids:
                files.append(UploadedFile.objects.get(id=fid))
        except UploadedFile.DoesNotExist:
            return redirect('converter:merge_pdf')

        # Perform PDF Merge
        from PyPDF2 import PdfMerger
        import os
        from django.conf import settings
        
        file_paths = [f.file.path for f in files]
        
        output_filename = f"merged_{uuid.uuid4()}.pdf"
        output_rel_path = f"processed/{output_filename}"
        output_full_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_full_path), exist_ok=True)
        
        try:
            merger = PdfMerger()
            for path in file_paths:
                merger.append(path)
            merger.write(output_full_path)
            merger.close()
            
            processed_file = ProcessedFile(file=output_rel_path)
            processed_file.save()
            
            # Track usage statistics
            try:
                from django.utils import timezone
                stat, _ = DailyStat.objects.get_or_create(date=timezone.now().date())
                stat.usage_count += 1
                stat.save()
            except Exception as e:
                print(f"Stats error: {e}")
            
            if 'merge_pdf_file_ids' in request.session:
                del request.session['merge_pdf_file_ids']
                
            return render(request, 'converter/result.html', {
                'download_url': processed_file.file.url,
                'file_id': processed_file.id,
                'file_type': 'PDF'
            })
            
        except Exception as e:
            print(f"PDF Merge Error: {e}")
            return redirect('converter:merge_pdf')

class ResultView(View):
    def get(self, request):
        job_id = request.GET.get('job_id')
        file_type = request.GET.get('file_type', 'WORD')
        
        if job_id:
            # Construct download URL for the result page button
            download_url = reverse('converter:download_file_direct', kwargs={'job_id': job_id})
            return render(request, 'converter/result.html', {
                'download_url': download_url,
                'file_type': file_type,
                'job_id': job_id
            })
            
        return render(request, 'converter/result.html')

from django.http import FileResponse, Http404
class DownloadFileView(View):
    def get(self, request, file_id):
        try:
            processed_file = ProcessedFile.objects.get(id=file_id)
            file_path = processed_file.file.path
            
            # Get custom filename from query params
            filename = request.GET.get('filename', '')
            if not filename:
                filename = processed_file.file.name.split('/')[-1]
            
            # Detect extension from original file
            original_ext = '.' + processed_file.file.name.split('.')[-1].lower()
            if not filename.lower().endswith(original_ext):
                # If user didn't type extension, append it. 
                # If user typed WRONG extension, maybe we should fix it? 
                # For now, just ensure it ends with the real extension to allow opening.
                # But wait, logic above was forcing .docx.
                # Let's align with the actual file type.
                if original_ext in ['.doc', '.docx']:
                     if not filename.lower().endswith('.docx') and not filename.lower().endswith('.doc'):
                          filename += '.docx'
                elif original_ext == '.pdf':
                     if not filename.lower().endswith('.pdf'):
                          filename += '.pdf'
                
            response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=filename)
            return response
        except ProcessedFile.DoesNotExist:
            raise Http404("File not found")
        except Exception as e:
            print(f"Download error: {e}")
            raise Http404("File processing error")


class TermsView(View):
    def get(self, request):
        return render(request, 'converter/terms.html', {
            'title': 'เงื่อนไขการใช้งาน',
            'subtitle': 'ข้อกำหนดและเงื่อนไขการใช้บริการของ MankyFile'
        })


class PrivacyView(View):
    def get(self, request):
        return render(request, 'converter/privacy.html', {
            'title': 'นโยบายความเป็นส่วนตัว',
            'subtitle': 'การคุ้มครองข้อมูลส่วนบุคคลและความเป็นส่วนตัวของผู้ใช้'
        })


# --- CUSTOM QR DRAWERS ---
from qrcode.image.styles.moduledrawers.pil import StyledPilQRModuleDrawer
from PIL import ImageDraw

class QRStyledEyeDrawer(StyledPilQRModuleDrawer):
    """
    Custom drawer for QR Corner Eyes (Finder Patterns)
    Supports different shapes for the frame and ball.
    """
    def __init__(self, style='square'):
        self.style = style
        self.drawn_eyes = set()

    def initialize(self, *args, **kwargs):
        super().initialize(*args, **kwargs)
        self.img_draw = ImageDraw.Draw(self.img._img)

    def drawrect(self, box, is_active):
        # Find module coordinates
        x, y = box[0]
        bs = self.img.box_size
        border = self.img.border
        col = int(round(x / bs)) - border
        row = int(round(y / bs)) - border
        width = self.img.width
        
        # Eyes are at 0..6 or (width-7)..(width-1)
        eye_positions = [(0, 0), (0, width - 7), (width - 7, 0)]
        
        for er, ec in eye_positions:
            if er <= row < er + 7 and ec <= col < ec + 7:
                # We only draw the entire eye once when we hit the first module of it
                if (er, ec) not in self.drawn_eyes:
                    self.draw_eye_shape(er, ec)
                    self.drawn_eyes.add((er, ec))
                return

    def draw_eye_shape(self, row, col):
        bs = self.img.box_size
        border = self.img.border
        # Use paint_color (maskable) and back_color (background)
        paint_color = self.img.paint_color
        back_color = self.img.color_mask.back_color

        x0 = (col + border) * bs
        y0 = (row + border) * bs
        x1 = x0 + (7 * bs)
        y1 = y0 + (7 * bs)

        # 1. Frame (Outer 7x7)
        self.draw_shape(x0, y0, x1, y1, paint_color, self.style)
        
        # 2. Gap (Clear inner 5x5)
        # Note: We use back_color to clear the gap
        g0, g1 = x0 + bs, x1 - bs
        h0, h1 = y0 + bs, y1 - bs
        self.draw_shape(g0, h0, g1, h1, back_color, self.style)

        # 3. Ball (Solid inner 3x3)
        b0, b1 = x0 + (2 * bs), x1 - (2 * bs)
        c0, c1 = y0 + (2 * bs), y1 - (2 * bs)
        self.draw_shape(b0, c0, b1, c1, paint_color, self.style)

    def draw_shape(self, x0, y0, x1, y1, color, style):
        # PIL handles coordinates as [left, top, right, bottom]
        rect = [x0, y0, x1-1, y1-1]
        
        if style == 'circle':
            self.img_draw.ellipse(rect, fill=color)
        elif style == 'rounded':
            radius = (x1 - x0) // 4
            self.img_draw.rounded_rectangle(rect, fill=color, radius=radius)
        elif style == 'leaf':
            # NW and SE rounded, others square
            # For simplicity, we use rounded with custom radii if supported, 
            # but standard PIL rounded_rectangle applies to all.
            # Fallback: just use rounded for now or custom implementation later
            radius = (x1 - x0) // 3
            self.img_draw.rounded_rectangle(rect, fill=color, radius=radius)
        elif style == 'shield':
            # Bottom corners rounded
            radius = (x1 - x0) // 3
            self.img_draw.rounded_rectangle(rect, fill=color, radius=radius)
        else: # square
            self.img_draw.rectangle(rect, fill=color)

class QRCodeGeneratorView(View):
    def get(self, request):
        from django.urls import reverse
        context = {
            'title': 'สร้าง QR Code',
            'subtitle': 'สร้าง QR Code ฟรีจากลิงก์ ข้อความ หรือเบอร์โทรศัพท์ พร้อมปรับแต่งสีได้ตามใจชอบ',
            'action_url': reverse('converter:qrcode_generator'),
        }
        return render(request, 'converter/qrcode_tool.html', context)

    def post(self, request):
        import qrcode
        from PIL import Image
        import os
        from django.conf import settings
        import uuid
        from django.urls import reverse
        from .models import UploadedFile, ProcessedFile, DailyStat
        from django.utils import timezone
        
        try:
            # 1. Get QR Type and Construct Data
            qr_type = request.POST.get('qr_type', 'url')
            data = ""
            
            if qr_type == 'url':
                data = request.POST.get('content_url', '').strip()
                if not data.startswith('http'): data = 'https://' + data
            
            elif qr_type == 'text':
                data = request.POST.get('content_text', '').strip()
                
            elif qr_type == 'wifi':
                ssid = request.POST.get('wifi_ssid', '').strip()
                password = request.POST.get('wifi_password', '').strip()
                encryption = request.POST.get('wifi_encryption', 'WPA')
                hidden = request.POST.get('wifi_hidden', 'false')
                # Format: WIFI:S:MySSID;T:WPA;P:MyPass;H:false;;
                data = f"WIFI:S:{ssid};T:{encryption};P:{password};H:{hidden};;"
            
            elif qr_type == 'email':
                email = request.POST.get('email_address', '').strip()
                subject = request.POST.get('email_subject', '').strip()
                body = request.POST.get('email_body', '').strip()
                data = f"mailto:{email}?subject={subject}&body={body}"
                
            elif qr_type == 'tel':
                phone = request.POST.get('tel_number', '').strip()
                data = f"tel:{phone}"
                
            elif qr_type == 'sms':
                phone = request.POST.get('sms_number', '').strip()
                message = request.POST.get('sms_message', '').strip()
                data = f"SMSTO:{phone}:{message}"
                
            elif qr_type == 'whatsapp':
                phone = request.POST.get('whatsapp_number', '').strip()
                message = request.POST.get('whatsapp_message', '').strip()
                # Clean phone number
                phone = ''.join(filter(str.isdigit, phone))
                data = f"https://wa.me/{phone}?text={message}"

            # Validate Data
            if not data:
                message = 'กรุณากรอกข้อมูลให้ครบถ้วน'
                if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                    return JsonResponse({'success': False, 'message': message})
                    
                context = {
                    'title': 'สร้าง QR Code',
                    'error': message,
                    'action_url': reverse('converter:qrcode_generator'),
                }
                return render(request, 'converter/qrcode_tool.html', context)
                
            # 2. Get Design Options
            fill_color = request.POST.get('fill_color', '#000000')
            back_color = request.POST.get('back_color', '#ffffff')
            pattern_style = request.POST.get('pattern_style', 'square')
            eye_style = request.POST.get('eye_style', 'square')
            
            # 3. Generate QR Code
            import qrcode
            from django.http import JsonResponse
            
            # Helper: Convert Hex to RGB Tuple
            def hex_to_rgb(hex_color):
                try:
                    hex_color = hex_color.lstrip('#')
                    if len(hex_color) == 3:
                        hex_color = ''.join([c*2 for c in hex_color])
                    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                except:
                    return (0, 0, 0) # Default black

            qr = qrcode.QRCode(
                version=None,
                error_correction=qrcode.constants.ERROR_CORRECT_H,
                box_size=15,
                border=4,
            )
            qr.add_data(data)
            qr.make(fit=True)
            
            img = None
            try:
                # qrcode 8.x+ structure
                from qrcode.image.styledpil import StyledPilImage
                from qrcode.image.styles.moduledrawers import (
                    SquareModuleDrawer, GappedSquareModuleDrawer, CircleModuleDrawer,
                    RoundedModuleDrawer, VerticalBarsDrawer, HorizontalBarsDrawer
                )
                from qrcode.image.styles.colormasks import SolidFillColorMask
                
                # Setup Module Drawer
                drawers = {
                    'square': SquareModuleDrawer(),
                    'gapped': GappedSquareModuleDrawer(),
                    'circle': CircleModuleDrawer(),
                    'rounded': RoundedModuleDrawer(),
                    'vertical': VerticalBarsDrawer(),
                    'horizontal': HorizontalBarsDrawer()
                }
                selected_drawer = drawers.get(pattern_style, SquareModuleDrawer())
                
                # Setup Eye Drawer
                selected_eye_drawer = QRStyledEyeDrawer(style=eye_style)
                
                # Construct Color Mask
                fill_rgb = hex_to_rgb(fill_color)
                back_rgb = hex_to_rgb(back_color)
                color_mask = SolidFillColorMask(back_color=back_rgb, front_color=fill_rgb)
                
                # Generate with StyledPilImage
                img = qr.make_image(
                    image_factory=StyledPilImage, 
                    module_drawer=selected_drawer,
                    eye_drawer=selected_eye_drawer,
                    color_mask=color_mask
                ).convert('RGB')
                
                print(f"[QR SUCCESS] Styled QR ({pattern_style}/{eye_style}) generated")
                
            except Exception as e:
                import traceback
                with open(os.path.join(settings.BASE_DIR, 'qr_error.log'), 'a', encoding='utf-8') as f:
                    f.write(f"\n--- {timezone.now()} ---\nStyle: {pattern_style}/{eye_style}\nError: {str(e)}\n")
                    f.write(traceback.format_exc())
                print(f"[QR STYLE ERROR] Falling back to standard: {e}")
                img = qr.make_image(fill_color=fill_color, back_color=back_color).convert('RGB')

            # --- LOGO PROCESSING ---
            logo_file = request.FILES.get('logo')
            if logo_file and img:
                try:
                    logo = Image.open(logo_file)
                    img_w, img_h = img.size
                    
                    # Calculate logo size (max 20% of QR)
                    logo_max_size = int(img_w * 0.2)
                    logo_w, logo_h = logo.size
                    
                    if logo_w > logo_h:
                        new_w = logo_max_size
                        new_h = int(logo_h * (logo_max_size / logo_w))
                    else:
                        new_h = logo_max_size
                        new_w = int(logo_w * (logo_max_size / logo_h))
                        
                    logo = logo.resize((new_w, new_h), Image.Resampling.LANCZOS)
                    
                    # Paste with transparency support
                    pos = ((img_w - new_w) // 2, (img_h - new_h) // 2)
                    if logo.mode == 'RGBA':
                        img.paste(logo, pos, logo)
                    else:
                        img.paste(logo, pos)
                except Exception as logo_err:
                    print(f"[QR LOGO ERROR] {logo_err}")
            
            # 4. Save file
            job_id = str(uuid.uuid4())
            output_dir = os.path.join(settings.MEDIA_ROOT, 'processed', job_id)
            os.makedirs(output_dir, exist_ok=True)
            
            output_filename = f"qr_{qr_type}_{job_id[:8]}.png"
            output_full_path = os.path.join(output_dir, output_filename)
            
            img.save(output_full_path, format="PNG")
            
            # Save Record
            processed_rel_path = f"processed/{job_id}/{output_filename}"
            processed_file = ProcessedFile(file=processed_rel_path)
            processed_file.save()
            
            # Track Usage
            try:
                stat, _ = DailyStat.objects.get_or_create(date=timezone.now().date())
                stat.usage_count += 1
                stat.save()
            except Exception as e:
                print(f"Stats error: {e}")
            
            # 5. Prepare Context/Response
            qr_image_url = f"{settings.MEDIA_URL}{processed_rel_path}"
            download_url = reverse('converter:download_file', kwargs={'file_id': processed_file.id})
            
            # Check for AJAX Request
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'qr_image_url': qr_image_url,
                    'download_url': download_url,
                    'message': 'QR Code generated successfully',
                    'debug_style': pattern_style
                })

            # Normal Request
            context = {
                'title': 'สร้าง QR Code',
                'subtitle': 'สร้าง QR Code ฟรีจากลิงก์ ข้อความ หรือเบอร์โทรศัพท์ พร้อมปรับแต่งสีได้ตามใจชอบ',
                'action_url': reverse('converter:qrcode_generator'),
                'generated': True,
                'qr_image_url': qr_image_url,
                'download_url': download_url,
                'fill_color': fill_color,
                'back_color': back_color,
                'pattern_style': pattern_style,
                'eye_style': eye_style,
            }

            return render(request, 'converter/qrcode_tool.html', context)
            
        except Exception as e:
            print(f"QR Error: {e}")
            import traceback
            traceback.print_exc()
            
            message = f'เกิดข้อผิดพลาด: {str(e)}'
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'message': message})
                
            context = {
                'title': 'สร้าง QR Code',
                'error': message,
                'action_url': reverse('converter:qrcode_generator'),
            }
            return render(request, 'converter/qrcode_tool.html', context)

# --- NEW IMAGE & PDF TOOLS (Worker Pattern) ---

class PDFToImageView(View):
    def get(self, request):
        context = {
            'title': 'PDF เป็น Image',
            'subtitle': 'แปลงหน้า PDF ของคุณเป็นไฟล์รูปภาพ (JPG/PNG) ได้อย่างง่ายดาย',
            'action_url': reverse('converter:pdf_to_image'),
            'file_type': 'PDF',
            'accept': '.pdf'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            elif 'file' in request.FILES:
                uploaded_file = request.FILES['file']
            else:
                return redirect('converter:index')
            
            upload_instance = UploadedFile.objects.create(file=uploaded_file, original_filename=uploaded_file.name)
            input_path = upload_instance.file.path
            
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            with open(input_path, 'rb') as f:
                files = {'file': (uploaded_file.name, f, 'application/pdf')}
                headers = {'X-API-KEY': API_KEY}
                data = {'type': 'pdf-to-image'} # ZIP of images
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    return render(request, 'converter/worker_wait.html', {
                        'task_id': task_id,
                        'worker_host': 'https://blilnkdex.biz.id',
                        'file_name': uploaded_file.name,
                        'task_type': 'ZIP' # Output usually ZIP for multi-page
                    })
                else:
                    raise Exception(f"Worker Error: {response.status_code}")
        except Exception as e:
            return render(request, 'converter/tool_base.html', {'error': str(e), 'title': 'PDF เป็น Image'})

class ImageToPDFView(View):
    def get(self, request):
        context = {
            'title': 'Image เป็น PDF',
            'subtitle': 'แปลงรูปภาพของคุณเป็นไฟล์ PDF ภายในไม่กี่วินาที',
            'action_url': reverse('converter:image_to_pdf'),
            'file_type': 'IMAGE',
            'accept': 'image/*'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        files = request.FILES.getlist('files')
        if files:
            uploaded_ids = []
            for f in files:
                uploaded_file = UploadedFile.objects.create(
                    file=f,
                    original_filename=f.name
                )
                uploaded_ids.append(str(uploaded_file.id))
            
            # Save IDs to session
            if request.GET.get('append') == 'true':
                 existing_ids = request.session.get('image_to_pdf_ids', [])
                 existing_ids.extend(uploaded_ids)
                 request.session['image_to_pdf_ids'] = existing_ids
                 request.session.modified = True
            else:
                 request.session['image_to_pdf_ids'] = uploaded_ids
                 request.session.modified = True
            
            return redirect('converter:arrange_image')
            
        return redirect('converter:image_to_pdf')

class ArrangeImageView(View):
    def get(self, request):
        file_ids = request.session.get('image_to_pdf_ids', [])
        if not file_ids:
            return redirect('converter:image_to_pdf')
        
        files = UploadedFile.objects.filter(id__in=file_ids)
        files_dict = {str(f.id): f for f in files}
        ordered_files = [files_dict[fid] for fid in file_ids if fid in files_dict]

        for f in ordered_files:
            f.display_name = f.original_filename if f.original_filename else f.filename

        context = {
            'files': ordered_files,
            'title': 'จัดเรียงรูปภาพ',
            'task_type': 'IMAGE_TO_PDF'
        }
        return render(request, 'converter/arrange_image.html', context)

    def post(self, request):
        ordered_ids = request.POST.getlist('file_order[]')
        
        if not ordered_ids:
            return redirect('converter:arrange_image')

        files = []
        try:
            for fid in ordered_ids:
                files.append(UploadedFile.objects.get(id=fid))
        except UploadedFile.DoesNotExist:
            return redirect('converter:image_to_pdf')

        import requests
        import zipfile
        import io
        
        try:
            # Re-order images and zip them for the worker
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                for i, f in enumerate(files):
                    zip_file.write(f.file.path, arcname=f"image_{i+1:03d}_{f.original_filename}")

            zip_buffer.seek(0)
            
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            # Send ZIP to worker with type image-to-pdf
            files_payload = {'file': ('images.zip', zip_buffer, 'application/zip')}
            headers = {'X-API-KEY': API_KEY}
            data = {'type': 'image-to-pdf'}
            
            target_url = f"{WORKER_URL}?key={API_KEY}"
            response = requests.post(target_url, files=files_payload, data=data, headers=headers, timeout=300, verify=False)
            
            if response.status_code == 200:
                res_data = response.json()
                task_id = res_data.get('task_id')
                
                # Cleanup session
                if 'image_to_pdf_ids' in request.session:
                    del request.session['image_to_pdf_ids']
                
                return render(request, 'converter/worker_wait.html', {
                    'task_id': task_id,
                    'worker_host': 'https://blilnkdex.biz.id',
                    'file_name': f"images_to_pdf_{task_id}.pdf",
                    'task_type': 'PDF'
                })
            else:
                raise Exception(f"Worker Error: {response.status_code}")
                
        except Exception as e:
            return render(request, 'converter/arrange_image.html', {
                'error': str(e),
                'files': files,
                'title': 'จัดเรียงรูปภาพ'
            })

class ImageResizeView(View):
    def get(self, request):
        context = {
            'title': 'ลดขนาดไฟล์ภาพ',
            'subtitle': 'ปรับขนาดหรือลดคุณภาพไฟล์ภาพเพื่อลดขนาดไฟล์',
            'action_url': reverse('converter:image_resize'),
            'file_type': 'IMAGE',
            'accept': 'image/*'
        }
        return render(request, 'converter/tool_base.html', context)

    def post(self, request):
        import requests
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            else:
                return redirect('converter:index')
            
            upload_instance = UploadedFile.objects.create(file=uploaded_file, original_filename=uploaded_file.name)
            input_path = upload_instance.file.path
            
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            with open(input_path, 'rb') as f:
                files = {'file': (uploaded_file.name, f, uploaded_file.content_type)}
                headers = {'X-API-KEY': API_KEY}
                data = {'type': 'image-resize'}
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    return render(request, 'converter/worker_wait.html', {
                        'task_id': task_id,
                        'worker_host': 'https://blilnkdex.biz.id',
                        'file_name': uploaded_file.name,
                        'task_type': 'Image'
                    })
                else:
                    raise Exception(f"Worker Error: {response.status_code}")
        except Exception as e:
            return render(request, 'converter/tool_base.html', {'error': str(e), 'title': 'ลดขนาดไฟล์ภาพ'})

class ImageConvertView(View):
    def get(self, request):
        context = {
            'title': 'แปลงไฟล์ภาพ',
            'subtitle': 'แปลงไฟล์รูปภาพเป็นนามสกุลอื่น เช่น JPG, PNG, WebP',
            'action_url': reverse('converter:image_convert'),
            'file_type': 'IMAGE',
            'accept': 'image/*'
        }
        return render(request, 'converter/image_convert_tool.html', context)

    def post(self, request):
        import requests
        try:
            if 'files' in request.FILES:
                uploaded_file = request.FILES.getlist('files')[0]
            else:
                return redirect('converter:index')
            
            target_format = request.POST.get('target_format', 'webp').upper() # JPG, PNG, WEBP
            
            upload_instance = UploadedFile.objects.create(file=uploaded_file, original_filename=uploaded_file.name)
            input_path = upload_instance.file.path
            
            WORKER_URL = 'https://blilnkdex.biz.id/api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            with open(input_path, 'rb') as f:
                files = {'file': (uploaded_file.name, f, uploaded_file.content_type)}
                headers = {'X-API-KEY': API_KEY}
                data = {
                    'type': 'image-convert',
                    'target_format': target_format.lower()
                }
                
                target_url = f"{WORKER_URL}?key={API_KEY}"
                response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
                
                if response.status_code == 200:
                    res_data = response.json()
                    task_id = res_data.get('task_id')
                    return render(request, 'converter/worker_wait.html', {
                        'task_id': task_id,
                        'worker_host': 'https://blilnkdex.biz.id',
                        'file_name': uploaded_file.name,
                        'task_type': target_format # Show correct type in result
                    })
                else:
                    raise Exception(f"Worker Error: {response.status_code}")
        except Exception as e:
            return render(request, 'converter/image_convert_tool.html', {'error': str(e), 'title': 'แปลงไฟล์ภาพ'})

class DeleteInstantView(View):
    def get(self, request):
        """
        Handle 'Delete Immediately' request.
        For now, this simply redirects to the index page, essentially 'discarding' the result.
        Actual file cleanup is handled by the background scheduler/cron job.
        """
        return redirect('converter:index')

# --- Shorten URL Views ---
class ShortenURLView(View):
    def get(self, request):
        return render(request, 'converter/shorten_link.html')
    
    def post(self, request):
        import requests
        original_url = request.POST.get('url')
        expiry_option = request.POST.get('expiry', '24h') 
        
        if not original_url:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.GET.get('ajax') == 'true':
                return JsonResponse({'success': False, 'error': 'กรุณาระบุ URL'})
            return render(request, 'converter/shorten_link.html', {'error': 'กรุณาระบุ URL'})
        
        try:
            WORKER_URL = 'https://blilnkdex.biz.id/shorten_api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            data = {'long_url': original_url, 'duration': expiry_option}
            headers = {'X-API-KEY': API_KEY}
            target_url = f'{WORKER_URL}?action=create&key={API_KEY}'
            response = requests.post(target_url, data=data, headers=headers, timeout=10, verify=False)
            
            if response.status_code == 200:
                res_data = response.json()
                if res_data.get('success'):
                    short_code = res_data.get('short_code')
                    short_url = request.build_absolute_uri(f'/s/{short_code}/')
                    expires_at_str = res_data.get('expires_at')
                    expires_at_display = expires_at_str if expires_at_str else 'ไม่มีวันหมดอายุ (Forever)'
                    
                    if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.GET.get('ajax') == 'true':
                        return JsonResponse({'success': True, 'short_url': short_url, 'expires_at': expires_at_display})

                    return render(request, 'converter/shorten_link.html', {
                        'short_url': short_url, 'original_url': original_url,
                        'expires_at': expires_at_display, 'short_code': short_code
                    })
                else:
                    raise Exception(res_data.get('error', 'Unknown Error'))
            else:
                raise Exception(f'Worker Error: {response.status_code}')
        except Exception as e:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.GET.get('ajax') == 'true':
                return JsonResponse({'success': False, 'error': str(e)})
            return render(request, 'converter/shorten_link.html', {'error': str(e)})

class RedirectShortLinkView(View):
    def get(self, request, short_code):
        return redirect(f'https://blilnkdex.biz.id/shorten_api.php?code={short_code}')

class SystemCleanupView(View):
    def get(self, request):
        # Security check (Simple key)
        key = request.GET.get('key')
        if key != settings.SECRET_KEY: # Use Django secret key or a specific one
            return JsonResponse({'status': 'error', 'message': 'Invalid Key'}, status=403)
            
        # Run Cleanup
        try:
            files_deleted = cleanup_old_files(hours=1)
            links_deleted = cleanup_expired_links_db()
            
            return JsonResponse({
                'status': 'success', 
                'files_deleted': files_deleted,
                'links_deleted': links_deleted,
                'message': 'System cleanup completed successfully'
            })
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

class HostImageView(View):
    def get(self, request):
        context = {
            'title': 'อัพโหลดรูปภาพ',
            'subtitle': 'ฝากรูปภาพออนไลน์พร้อมระบบลบอัตโนมัติ ใช้งานง่าย ปลอดภัย',
            'file_type': 'IMAGE',
            'accept': 'image/*'
        }
        return render(request, 'converter/host_image.html', context)

    def post(self, request):
        import requests
        try:
            image = request.FILES.get('image')
            auto_delete = request.POST.get('auto_delete', 'never')
            
            if not image:
                return redirect('converter:host_image')

            # Worker Configuration
            WORKER_URL = 'https://blilnkdex.biz.id/image_api.php'
            API_KEY = 'MANKY_SECRET_KEY_12345'
            
            data = {'auto_delete': auto_delete}
            headers = {'X-API-KEY': API_KEY}
            
            target_url = f'{WORKER_URL}?action=upload&key={API_KEY}'
            
            # Ensure we are reading from the start of the file
            image.seek(0)
            file_content = image.read()
            print(f"DEBUG: Uploading image to Worker using {target_url} | Size: {len(file_content)} bytes")
            
            files = {'image': (image.name, file_content, image.content_type)}
            response = requests.post(target_url, files=files, data=data, headers=headers, timeout=300, verify=False)
            
            if response.status_code == 200:
                res_data = response.json()
                if res_data.get('success'):
                    return render(request, 'converter/host_image_result.html', {
                        'res': res_data,
                        'title': 'ผลการอัพโหลดรูปภาพ'
                    })
                else:
                    raise Exception(res_data.get('error', 'Unknown error'))
            else:
                raise Exception(f'Worker Error: {response.status_code}')

        except Exception as e:
            return render(request, 'converter/host_image.html', {
                'error': str(e),
                'title': 'อัพโหลดรูปภาพ'
            })
